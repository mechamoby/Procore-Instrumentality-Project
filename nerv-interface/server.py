#!/usr/bin/env python3
"""NERV Interface v0.3 — Full agent session access via Gateway WebSocket.
Uses chat.send for real agent runs (tools, files, commands, memory).
"""

import logging
import asyncio
import json
import os
import re
import uuid
import time
from pathlib import Path

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File
from eva_sentry import EVASentry
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
import httpx
import uvicorn
import websockets

# Config
GATEWAY_HOST = "127.0.0.1"
GATEWAY_PORT = 18789
GATEWAY_TOKEN = os.environ.get("OPENCLAW_GATEWAY_TOKEN", "")
TTYD_HOST = "192.168.8.124"
TTYD_PORT = 7681
NERV_PORT = 8080
NERV_SESSION_KEY = "agent:main:nerv"

app = FastAPI(title="NERV Interface")

STATIC_DIR = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

# Serve index.html with no-cache headers to prevent stale versions
from fastapi.responses import HTMLResponse

@app.get("/")
async def serve_index():
    index_path = STATIC_DIR / "index.html"
    content = index_path.read_text()
    return HTMLResponse(content=content, headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0",
    })

UPLOAD_DIR = Path.home() / ".openclaw" / "workspace" / "nerv-uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
SENTRY_STATE_DIR = Path.home() / ".openclaw" / "workspace" / "eva-sentry-v1" / "state"
SENTRY = EVASentry(SENTRY_STATE_DIR)


@app.get("/config")
async def config():
    return {
        "ttyd_url": f"http://{TTYD_HOST}:{TTYD_PORT}",
        "ws_url": f"ws://{{host}}:{NERV_PORT}/ws/chat",
    }


async def create_gateway_ws():
    """Create an authenticated Gateway WebSocket connection."""
    headers = {"Origin": "http://127.0.0.1:8080"}
    ws = await websockets.connect(
        f"ws://{GATEWAY_HOST}:{GATEWAY_PORT}",
        max_size=10 * 1024 * 1024,
        additional_headers=headers,
    )
    
    # Receive challenge
    await asyncio.wait_for(ws.recv(), timeout=5)
    
    # Connect handshake
    connect_msg = {
        "type": "req",
        "id": "connect-1",
        "method": "connect",
        "params": {
            "minProtocol": 3,
            "maxProtocol": 3,
            "client": {
                "id": "openclaw-control-ui",
                "version": "0.3",
                "platform": "linux",
                "mode": "webchat",
            },
            "role": "operator",
            "scopes": ["operator.read", "operator.write", "operator.admin"],
            "caps": [],
            "commands": [],
            "permissions": {},
            "auth": {"token": GATEWAY_TOKEN},
            "locale": "en-US",
            "userAgent": "nerv-interface/0.3",
        },
    }
    await ws.send(json.dumps(connect_msg))
    resp = json.loads(await asyncio.wait_for(ws.recv(), timeout=10))
    
    if not resp.get("ok"):
        raise Exception(f"Gateway connect failed: {resp.get('error')}")
    
    return ws


@app.websocket("/ws/chat")
async def chat_ws(websocket: WebSocket):
    """WebSocket bridge: NERV client <-> OpenClaw Gateway agent session."""
    await websocket.accept()
    
    gw_ws = None
    try:
        # Connect to gateway
        try:
            gw_ws = await create_gateway_ws()
        except Exception as e:
            await websocket.send_text(json.dumps({
                "type": "error",
                "content": f"Cannot connect to OpenClaw gateway: {e}"
            }))
            return
        
        await websocket.send_text(json.dumps({
            "type": "system",
            "content": "Connected to OpenClaw Gateway (full agent access)"
        }))
        
        # Background task: listen to gateway events and forward to client
        current_run_id = None
        run_complete = asyncio.Event()
        last_chunk = None  # Dedup: track last sent chunk to avoid doubles
        
        async def gateway_listener():
            nonlocal current_run_id
            try:
                async for raw in gw_ws:
                    try:
                        msg = json.loads(raw)
                    except json.JSONDecodeError:
                        continue
                    
                    msg_type = msg.get("type")
                    
                    if msg_type == "event":
                        event_name = msg.get("event", "")
                        payload = msg.get("payload", {})
                        
                        if not isinstance(payload, dict):
                            continue
                        
                        session_key = payload.get("sessionKey", "")
                        run_id = payload.get("runId", "")
                        
                        # Only process events for our session
                        if session_key and session_key != NERV_SESSION_KEY:
                            continue
                        
                        if event_name == "agent":
                            stream = payload.get("stream", "")
                            data = payload.get("data", {})
                            
                            if stream == "assistant":
                                delta = data.get("delta", "")
                                if delta:
                                    await websocket.send_text(json.dumps({
                                        "type": "chunk",
                                        "content": delta
                                    }))
                            
                            elif stream == "tool":
                                tool_name = data.get("name", "")
                                phase = data.get("phase", "")
                                
                                if phase == "start":
                                    await websocket.send_text(json.dumps({
                                        "type": "tool_start",
                                        "tool": tool_name,
                                        "input": data.get("input", {})
                                    }))
                                elif phase == "end":
                                    await websocket.send_text(json.dumps({
                                        "type": "tool_end",
                                        "tool": tool_name,
                                        "status": data.get("status", "ok")
                                    }))
                            
                            elif stream == "lifecycle":
                                phase = data.get("phase", "")
                                if phase == "start":
                                    await websocket.send_text(json.dumps({
                                        "type": "typing"
                                    }))
                                elif phase == "end":
                                    await websocket.send_text(json.dumps({
                                        "type": "done"
                                    }))
                                    current_run_id = None
                                    run_complete.set()
                        
                        elif event_name == "chat":
                            state = payload.get("state", "")
                            
                            if state == "delta":
                                message = payload.get("message", {})
                                content = message.get("content", [])
                                for block in (content if isinstance(content, list) else []):
                                    if block.get("type") == "text":
                                        # Already handled by agent stream
                                        pass
                            
                            elif state == "final":
                                # Final message — ensure done is sent
                                if current_run_id:
                                    await websocket.send_text(json.dumps({
                                        "type": "done"
                                    }))
                                    current_run_id = None
                                    run_complete.set()
                    
                    elif msg_type == "res":
                        # Handle response to chat.send
                        payload = msg.get("payload", {})
                        if isinstance(payload, dict):
                            status = payload.get("status", "")
                            if status == "started":
                                current_run_id = payload.get("runId")
                            elif status == "error":
                                await websocket.send_text(json.dumps({
                                    "type": "error",
                                    "content": payload.get("error", "Run failed")
                                }))
                                run_complete.set()
                        
                        if not msg.get("ok") and msg.get("error"):
                            err = msg.get("error", {})
                            err_msg = err.get("message", str(err)) if isinstance(err, dict) else str(err)
                            await websocket.send_text(json.dumps({
                                "type": "error",
                                "content": err_msg
                            }))
                            run_complete.set()
            
            except websockets.exceptions.ConnectionClosed:
                await websocket.send_text(json.dumps({
                    "type": "error",
                    "content": "Gateway connection lost"
                }))
            except Exception as e:
                await websocket.send_text(json.dumps({
                    "type": "error",
                    "content": f"Gateway listener error: {e}"
                }))
        
        # Start gateway listener
        listener_task = asyncio.create_task(gateway_listener())
        
        # Track if this is first message in session (needs boot context)
        is_first_message = True
        
        # Main loop: receive from NERV client, forward to gateway
        try:
            while True:
                data = await websocket.receive_text()
                msg = json.loads(data)
                user_text = msg.get("content", "")
                attachments = msg.get("attachments", [])
                
                if not user_text and not attachments:
                    continue
                
                run_complete.clear()
                
                # First message flag (reserved for future use)
                if is_first_message:
                    is_first_message = False
                
                # Build message text — if files attached, prepend file context
                message_text = user_text or ""
                gateway_attachments = []
                
                if attachments:
                    file_lines = []
                    for att in attachments:
                        media_path = att.get("mediaPath", att.get("path", ""))
                        filename = att.get("filename", "file")
                        mime = att.get("mime", "application/octet-stream")
                        data_url = att.get("dataUrl", "")
                        
                        # Add file path context so agent knows where the file is on disk
                        file_lines.append(
                            f"[media attached: {media_path} ({mime}) | {media_path}]"
                        )
                        
                        # Build gateway attachment from base64 dataURL
                        # Format matches OpenClaw Control UI: {type, mimeType, content}
                        if data_url:
                            m = re.match(r'^data:([^;]+);base64,(.+)$', data_url)
                            if m:
                                gateway_attachments.append({
                                    "type": "image",
                                    "mimeType": m.group(1),
                                    "content": m.group(2),
                                })
                    
                    file_context = "\n".join(file_lines)
                    if message_text:
                        message_text = f"{file_context}\n{message_text}"
                    else:
                        message_text = file_context
                
                # EVA Sentry preflight: prompt-injection / malware intent in text
                text_scan = SENTRY.scan_text(message_text or "")
                if text_scan.get("verdict") in {"deny", "quarantine"}:
                    await websocket.send_text(json.dumps({
                        "type": "error",
                        "content": "EVA Sentry blocked message (prompt-injection/malware pattern)."
                    }))
                    continue

                # EVA Sentry preflight: attached file verdict enforcement
                blocked_attachment = None
                for att in attachments or []:
                    p = att.get("mediaPath") or att.get("path")
                    if not p:
                        continue
                    v = SENTRY.get_verdict_for_path(p) or SENTRY.get_verdict_for_path(att.get("path", ""))
                    if not v:
                        # If attachment somehow skipped upload scan, fail closed.
                        blocked_attachment = {"path": p, "reason": "missing_verdict"}
                        break
                    if v.get("verdict") in {"quarantine", "deny"}:
                        blocked_attachment = {"path": p, "reason": ", ".join(v.get("reasons", []))}
                        break

                if blocked_attachment:
                    await websocket.send_text(json.dumps({
                        "type": "error",
                        "content": f"EVA Sentry blocked attachment: {blocked_attachment['path']} ({blocked_attachment['reason']})"
                    }))
                    continue

                # Send via chat.send for full agent session
                idempotency_key = f"nerv-{uuid.uuid4().hex[:12]}"
                chat_params = {
                    "sessionKey": NERV_SESSION_KEY,
                    "message": message_text,
                    "deliver": False,
                    "idempotencyKey": idempotency_key,
                }
                if gateway_attachments:
                    chat_params["attachments"] = gateway_attachments

                chat_req = {
                    "type": "req",
                    "id": idempotency_key,
                    "method": "chat.send",
                    "params": chat_params,
                }
                await gw_ws.send(json.dumps(chat_req))
        
        except WebSocketDisconnect:
            pass
        finally:
            listener_task.cancel()
    
    finally:
        if gw_ws:
            await gw_ws.close()


@app.get("/api/status")
async def get_status():
    async with httpx.AsyncClient(timeout=10) as client:
        try:
            r = await client.post(
                f"http://{GATEWAY_HOST}:{GATEWAY_PORT}/tools/invoke",
                headers={
                    "Authorization": f"Bearer {GATEWAY_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={"tool": "session_status", "args": {}},
            )
            return r.json()
        except Exception as e:
            return {"error": str(e)}


@app.get("/api/sessions")
async def get_sessions():
    async with httpx.AsyncClient(timeout=10) as client:
        try:
            r = await client.post(
                f"http://{GATEWAY_HOST}:{GATEWAY_PORT}/tools/invoke",
                headers={
                    "Authorization": f"Bearer {GATEWAY_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={"tool": "sessions_list", "args": {"limit": 10, "messageLimit": 1}},
            )
            return r.json()
        except Exception as e:
            return {"error": str(e)}


@app.get("/api/subagents")
async def get_subagents():
    """Get sub-agent status directly from the OpenClaw subagents tool."""
    async with httpx.AsyncClient(timeout=15) as client:
        try:
            r = await client.post(
                f"http://{GATEWAY_HOST}:{GATEWAY_PORT}/tools/invoke",
                headers={
                    "Authorization": f"Bearer {GATEWAY_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={"tool": "subagents", "args": {"action": "list", "recentMinutes": 1440}},
            )
            data = r.json()
            # Pass through normalized tool output for UI compatibility
            return {
                "active": data.get("active", []),
                "recent": data.get("recent", []),
                "total": data.get("total", 0),
            }
        except Exception as e:
            return {"error": str(e)}


@app.post("/api/subagents/steer")
async def steer_subagent(data: dict):
    """Send a steer message to a sub-agent."""
    target = data.get("target", "")
    message = data.get("message", "")
    if not target or not message:
        return {"error": "target and message required"}
    async with httpx.AsyncClient(timeout=15) as client:
        try:
            r = await client.post(
                f"http://{GATEWAY_HOST}:{GATEWAY_PORT}/tools/invoke",
                headers={
                    "Authorization": f"Bearer {GATEWAY_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={"tool": "subagents", "args": {"action": "steer", "target": target, "message": message}},
            )
            return r.json()
        except Exception as e:
            return {"error": str(e)}


@app.post("/api/subagents/kill")
async def kill_subagent(data: dict):
    """Kill a sub-agent."""
    target = data.get("target", "")
    if not target:
        return {"error": "target required"}
    async with httpx.AsyncClient(timeout=15) as client:
        try:
            r = await client.post(
                f"http://{GATEWAY_HOST}:{GATEWAY_PORT}/tools/invoke",
                headers={
                    "Authorization": f"Bearer {GATEWAY_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={"tool": "subagents", "args": {"action": "kill", "target": target}},
            )
            return r.json()
        except Exception as e:
            return {"error": str(e)}


@app.get("/api/system")
async def get_system_metrics():
    """Get system metrics.

    Primary path uses psutil. If unavailable, falls back to /proc + os.statvfs
    so the NERV dashboard still renders live values instead of hard errors.
    """
    try:
        try:
            import psutil

            cpu = psutil.cpu_percent(interval=0.25)
            mem = psutil.virtual_memory()
            disk = psutil.disk_usage('/')
            boot = psutil.boot_time()
            ram_pct = mem.percent
            disk_pct = disk.percent
            ram_used = mem.used
            ram_total = mem.total
            disk_used = disk.used
            disk_total = disk.total
        except Exception:
            # Fallback path without psutil
            def _read_first(path: str) -> str:
                with open(path, 'r', encoding='utf-8') as f:
                    return f.readline().strip()

            # CPU fallback from /proc/loadavg (1m load normalized to cores)
            load1 = float(_read_first('/proc/loadavg').split()[0])
            cores = os.cpu_count() or 1
            cpu = max(0.0, min(100.0, (load1 / cores) * 100.0))

            # Memory fallback from /proc/meminfo
            meminfo = {}
            with open('/proc/meminfo', 'r', encoding='utf-8') as f:
                for line in f:
                    k, v = line.split(':', 1)
                    meminfo[k.strip()] = int(v.strip().split()[0]) * 1024
            ram_total = meminfo.get('MemTotal', 0)
            mem_available = meminfo.get('MemAvailable', 0)
            ram_used = max(0, ram_total - mem_available)
            ram_pct = (ram_used / ram_total * 100.0) if ram_total else 0.0

            # Disk fallback from statvfs
            st = os.statvfs('/')
            disk_total = st.f_blocks * st.f_frsize
            disk_free = st.f_bavail * st.f_frsize
            disk_used = max(0, disk_total - disk_free)
            disk_pct = (disk_used / disk_total * 100.0) if disk_total else 0.0

            # Uptime fallback from /proc/uptime
            uptime_raw = float(_read_first('/proc/uptime').split()[0])
            boot = time.time() - uptime_raw

        # Uptime
        uptime_secs = int(time.time() - boot)
        days, rem = divmod(uptime_secs, 86400)
        hours, rem = divmod(rem, 3600)
        mins, _ = divmod(rem, 60)
        if days > 0:
            uptime_str = f"{days}d {hours}h {mins}m"
        elif hours > 0:
            uptime_str = f"{hours}h {mins}m"
        else:
            uptime_str = f"{mins}m"

        metrics = {
            "cpu": f"{cpu:.1f}%",
            "ram": f"{ram_used // (1024**3)}GB / {ram_total // (1024**3)}GB",
            "disk": f"{disk_used // (1024**3)}GB / {disk_total // (1024**3)}GB",
            "uptime": uptime_str,
            "cpu_pct": cpu,
            "ram_pct": ram_pct,
            "disk_pct": disk_pct,
        }
        return {"metrics": metrics}
    except Exception as e:
        return {"error": str(e)}


@app.post("/api/switch-session")
async def switch_session(data: dict):
    """Switch the active session key."""
    global NERV_SESSION_KEY
    new_key = data.get("sessionKey")
    if not new_key:
        return {"error": "sessionKey required"}
    
    NERV_SESSION_KEY = new_key
    return {"success": True, "sessionKey": NERV_SESSION_KEY}


@app.get("/api/current-session")
async def get_current_session():
    """Get the current active session key."""
    return {"sessionKey": NERV_SESSION_KEY}


@app.get("/api/files")
async def list_files(path: str = ""):
    """List directory contents."""
    workspace = Path.home() / ".openclaw" / "workspace"
    
    if not path:
        target = workspace
    else:
        target = workspace / path.lstrip("/")
    
    # Security check - ensure we stay within workspace
    try:
        target = target.resolve()
        workspace = workspace.resolve()
        if not str(target).startswith(str(workspace)):
            return {"error": "Access denied"}
    except Exception:
        return {"error": "Invalid path"}
    
    if not target.exists():
        return {"error": "Path not found"}
    
    if not target.is_dir():
        return {"error": "Not a directory"}
    
    try:
        items = []
        for item in sorted(target.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower())):
            rel_path = str(item.relative_to(workspace))
            items.append({
                "name": item.name,
                "path": rel_path,
                "isDirectory": item.is_dir(),
                "size": item.stat().st_size if item.is_file() else None,
                "modified": item.stat().st_mtime
            })
        
        return {"items": items, "currentPath": str(target.relative_to(workspace)) if target != workspace else ""}
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/files/read")
async def read_file(path: str):
    """Read file content."""
    workspace = Path.home() / ".openclaw" / "workspace"
    target = workspace / path.lstrip("/")
    
    # Security check
    try:
        target = target.resolve()
        workspace = workspace.resolve()
        if not str(target).startswith(str(workspace)):
            return {"error": "Access denied"}
    except Exception:
        return {"error": "Invalid path"}
    
    if not target.exists():
        return {"error": "File not found"}
    
    if not target.is_file():
        return {"error": "Not a file"}
    
    try:
        # Check if file is too large
        if target.stat().st_size > 1024 * 1024:  # 1MB limit
            return {"error": "File too large (max 1MB)"}
        
        with open(target, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "content": content,
            "path": path,
            "size": target.stat().st_size,
            "modified": target.stat().st_mtime
        }
    except UnicodeDecodeError:
        return {"error": "Binary file - cannot display"}
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/file")
async def serve_file(path: str):
    """Serve a file for download/viewing."""
    workspace = Path.home() / ".openclaw" / "workspace"
    target = workspace / path.lstrip("/")
    
    # Security check
    try:
        target = target.resolve()
        workspace = workspace.resolve()
        if not str(target).startswith(str(workspace)):
            return {"error": "Access denied"}
    except Exception:
        return {"error": "Invalid path"}
    
    if not target.exists() or not target.is_file():
        return {"error": "File not found"}
    
    return FileResponse(str(target))


@app.get("/api/media")
async def serve_media(path: str):
    """Serve uploaded files from nerv-uploads or media/inbound."""
    target = Path(path)
    # Allow serving from upload dir and media inbound
    allowed = [
        UPLOAD_DIR.resolve(),
        (Path.home() / ".openclaw" / "media" / "inbound").resolve(),
    ]
    try:
        resolved = target.resolve()
        if not any(str(resolved).startswith(str(a)) for a in allowed):
            return {"error": "Access denied"}
    except Exception:
        return {"error": "Invalid path"}
    if not resolved.exists() or not resolved.is_file():
        return {"error": "File not found"}
    return FileResponse(str(resolved))


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload a file and return its path. The NERV frontend can then reference
    it in a chat message so the agent can read/process it."""
    import shutil
    safe_name = file.filename.replace("/", "_").replace("\\", "_")
    ts = int(time.time())
    dest = UPLOAD_DIR / f"{ts}_{safe_name}"
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Also copy to the inbound media dir so the agent's media pipeline can find it
    media_dir = Path.home() / ".openclaw" / "media" / "inbound"
    media_dir.mkdir(parents=True, exist_ok=True)
    media_dest = media_dir / f"nerv_{ts}_{safe_name}"
    import shutil as sh
    sh.copy2(dest, media_dest)

    # EVA Sentry ingest scan (document malware + prompt-injection payload checks)
    sentry_verdict = SENTRY.scan_file(dest, declared_mime=file.content_type or "")
    # Mirror verdict for media path key lookups
    SENTRY.save_verdict_dict(str(media_dest), sentry_verdict)

    return {
        "path": str(dest),
        "mediaPath": str(media_dest),
        "filename": safe_name,
        "size": dest.stat().st_size,
        "mime": file.content_type or "application/octet-stream",
        "sentry": sentry_verdict,
    }


@app.post("/api/sentry/scan")
async def sentry_scan(data: dict):
    """Scan an existing file path (used by external ingest pipelines like email/Procore)."""
    target = (data or {}).get("path", "")
    mime = (data or {}).get("mime", "")
    if not target:
        return {"error": "path required"}
    p = Path(target)
    verdict = SENTRY.scan_file(p, declared_mime=mime)
    return {"ok": True, "path": str(p), "sentry": verdict}


@app.get("/api/sentry/verdict")
async def sentry_verdict(path: str):
    if not path:
        return {"error": "path required"}
    verdict = SENTRY.get_verdict_for_path(path)
    if not verdict:
        return {"error": "verdict not found"}
    return {"ok": True, "path": path, "sentry": verdict}


DOCUMENTS_DIR = Path(__file__).parent / "documents"
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)


@app.get("/api/documents")
async def list_documents():
    """List all saved documents."""
    docs = []
    for f in sorted(DOCUMENTS_DIR.glob("*.html"), key=lambda p: p.stat().st_mtime, reverse=True):
        stat = f.stat()
        docs.append({
            "name": f.stem,
            "size": stat.st_size,
            "modified": stat.st_mtime,
        })
    return {"documents": docs}


@app.post("/api/documents/save")
async def save_document(data: dict):
    """Save document as HTML and optionally convert to DOCX."""
    name = data.get("name", "Untitled").strip()
    html = data.get("html", "")
    fmt = data.get("format", "")

    if not name:
        return {"error": "name required"}

    safe_name = re.sub(r'[^\w\s\-.]', '', name).strip()
    if not safe_name:
        safe_name = "Untitled"

    # EVA Sentry scan for injection/malware-like payloads in document body
    doc_scan = SENTRY.scan_text(html or "")
    if doc_scan.get("verdict") in {"deny", "quarantine"}:
        return {"error": "EVA Sentry blocked document content", "sentry": doc_scan}

    # Save HTML
    html_path = DOCUMENTS_DIR / f"{safe_name}.html"
    html_path.write_text(html, encoding="utf-8")

    # Convert to DOCX if requested
    docx_path = None
    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            import html as html_mod
            from html.parser import HTMLParser

            class _DocxBuilder(HTMLParser):
                def __init__(self, doc):
                    super().__init__()
                    self.doc = doc
                    self.current_para = None
                    self.bold = False
                    self.italic = False
                    self.underline = False
                    self.in_heading = 0
                    self.in_list = None  # 'ul' or 'ol'
                    self.list_counter = 0
                    self.in_blockquote = False
                    self.text_buf = ""

                def _flush(self):
                    if self.text_buf and self.current_para is not None:
                        run = self.current_para.add_run(self.text_buf)
                        run.bold = self.bold
                        run.italic = self.italic
                        run.underline = self.underline
                        self.text_buf = ""

                def handle_starttag(self, tag, attrs):
                    tag = tag.lower()
                    if tag in ('h1', 'h2', 'h3'):
                        self._flush()
                        level = int(tag[1])
                        self.current_para = self.doc.add_heading('', level=level)
                        self.in_heading = level
                    elif tag == 'p':
                        self._flush()
                        if self.in_blockquote:
                            self.current_para = self.doc.add_paragraph()
                            self.current_para.paragraph_format.left_indent = Inches(0.5)
                        else:
                            self.current_para = self.doc.add_paragraph()
                    elif tag == 'strong' or tag == 'b':
                        self._flush()
                        self.bold = True
                    elif tag == 'em' or tag == 'i':
                        self._flush()
                        self.italic = True
                    elif tag == 'u':
                        self._flush()
                        self.underline = True
                    elif tag == 'ul':
                        self.in_list = 'ul'
                        self.list_counter = 0
                    elif tag == 'ol':
                        self.in_list = 'ol'
                        self.list_counter = 0
                    elif tag == 'li':
                        self._flush()
                        self.list_counter += 1
                        self.current_para = self.doc.add_paragraph()
                        if self.in_list == 'ol':
                            self.current_para.style = 'List Number'
                        else:
                            self.current_para.style = 'List Bullet'
                    elif tag == 'blockquote':
                        self.in_blockquote = True
                    elif tag == 'br':
                        self.text_buf += "\n"

                def handle_endtag(self, tag):
                    tag = tag.lower()
                    if tag in ('h1', 'h2', 'h3'):
                        self._flush()
                        self.in_heading = 0
                        self.current_para = None
                    elif tag == 'p':
                        self._flush()
                        self.current_para = None
                    elif tag == 'strong' or tag == 'b':
                        self._flush()
                        self.bold = False
                    elif tag == 'em' or tag == 'i':
                        self._flush()
                        self.italic = False
                    elif tag == 'u':
                        self._flush()
                        self.underline = False
                    elif tag in ('ul', 'ol'):
                        self._flush()
                        self.in_list = None
                    elif tag == 'li':
                        self._flush()
                        self.current_para = None
                    elif tag == 'blockquote':
                        self._flush()
                        self.in_blockquote = False

                def handle_data(self, data):
                    self.text_buf += data

            doc = DocxDocument()
            builder = _DocxBuilder(doc)
            builder.feed(html)
            builder._flush()

            docx_path = DOCUMENTS_DIR / f"{safe_name}.docx"
            doc.save(str(docx_path))

        except ImportError:
            return {"ok": True, "html_path": str(html_path), "docx_error": "python-docx not installed"}
        except Exception as e:
            return {"ok": True, "html_path": str(html_path), "docx_error": str(e)}

    return {"ok": True, "html_path": str(html_path), "docx_path": str(docx_path) if docx_path else None, "sentry": doc_scan}


@app.get("/api/documents/load")
async def load_document(name: str):
    """Load a document's HTML content."""
    safe_name = re.sub(r'[^\w\s\-.]', '', name).strip()
    html_path = DOCUMENTS_DIR / f"{safe_name}.html"
    if not html_path.exists():
        return {"error": "Document not found"}
    return {"name": safe_name, "html": html_path.read_text(encoding="utf-8")}


@app.get("/api/documents/export")
async def export_document(name: str, fmt: str = "docx"):
    """Export a saved document as DOCX or HTML download."""
    safe_name = re.sub(r'[^\w\s\-.]', '', name).strip()
    fmt = (fmt or "docx").lower()
    if fmt not in ("docx", "html"):
        return {"error": "Unsupported format"}

    file_path = DOCUMENTS_DIR / f"{safe_name}.{fmt}"
    if not file_path.exists():
        return {"error": f"Document '{safe_name}.{fmt}' not found"}

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if fmt == "docx" else "text/html"
    return FileResponse(path=file_path, media_type=media_type, filename=file_path.name)


@app.delete("/api/documents/delete")
async def delete_document(name: str):
    """Delete a document."""
    safe_name = re.sub(r'[^\w\s\-.]', '', name).strip()
    for ext in ('.html', '.docx'):
        p = DOCUMENTS_DIR / f"{safe_name}{ext}"
        if p.exists():
            p.unlink()
    return {"ok": True}


@app.post("/api/documents/email")
async def email_document(request_data: dict = None):
    """Email a document as .docx attachment."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email.mime.text import MIMEText
    from email import encoders
    
    if not request_data:
        return {"error": "No data provided"}
    
    name = request_data.get("name", "Untitled")
    to_addr = request_data.get("to", "")
    
    if not to_addr or "@" not in to_addr:
        return {"error": "Invalid email address"}
    
    safe_name = re.sub(r'[^\w\s\-.]', '', name).strip()
    docx_path = DOCUMENTS_DIR / f"{safe_name}.docx"
    
    if not docx_path.exists():
        return {"error": f"Document '{safe_name}.docx' not found. Save it first."}
    
    # Load SMTP credentials
    creds_path = Path.home() / ".credentials" / "smtp.env"
    if not creds_path.exists():
        # Try gmail creds
        creds_path = Path.home() / ".credentials" / "gmail.env"
    
    smtp_host = os.environ.get("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")
    
    # Try loading from credential files
    if not smtp_user:
        for env_file in [Path.home() / ".credentials" / "smtp.env", Path.home() / ".credentials" / "gmail.env"]:
            if env_file.exists():
                for line in env_file.read_text().splitlines():
                    line = line.strip()
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        k = k.strip().upper()
                        v = v.strip().strip('"').strip("'")
                        if k in ("SMTP_USER", "GMAIL_USER", "EMAIL_USER", "EMAIL"):
                            smtp_user = v
                        elif k in ("SMTP_PASS", "GMAIL_PASS", "EMAIL_PASS", "APP_PASSWORD"):
                            smtp_pass = v
                        elif k in ("SMTP_HOST",):
                            smtp_host = v
                        elif k in ("SMTP_PORT",):
                            smtp_port = int(v)
                break
    
    if not smtp_user or not smtp_pass:
        return {"error": "SMTP credentials not configured. Add ~/.credentials/smtp.env with SMTP_USER and SMTP_PASS."}
    
    try:
        msg = MIMEMultipart()
        msg["From"] = smtp_user
        msg["To"] = to_addr
        msg["Subject"] = f"{safe_name} — Sent from NERV"
        
        body = f"Document '{safe_name}' attached.\n\nSent from NERV Command Interface."
        msg.attach(MIMEText(body, "plain"))
        
        with open(docx_path, "rb") as f:
            part = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={safe_name}.docx")
            msg.attach(part)
        
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)
        
        return {"ok": True, "message": f"Sent to {to_addr}"}
    
    except Exception as e:
        return {"error": f"Email failed: {str(e)}"}


@app.websocket("/ws/activity")
async def activity_ws(websocket: WebSocket):
    """Stream live agent activity from ALL active session logs."""
    await websocket.accept()
    session_dir = Path.home() / ".openclaw" / "agents" / "main" / "sessions"
    
    try:
        # Track file positions for each session file we're tailing
        file_positions = {}  # path -> last read position
        
        # On connect, replay last 50 lines from the most recent session
        logs = sorted(
            [p for p in session_dir.glob("*.jsonl") if ".deleted" not in p.name and ".lock" not in p.name],
            key=lambda p: p.stat().st_mtime, reverse=True
        )
        if logs:
            # Replay from most recent session
            with open(logs[0], "r") as f:
                all_lines = f.readlines()
                replay_lines = all_lines[-50:] if len(all_lines) > 50 else all_lines
                for line in replay_lines:
                    await _process_activity_line(line, websocket)
            # Start tailing from current end of all recent sessions
            for log_file in logs[:5]:  # Track up to 5 most recent sessions
                file_positions[str(log_file)] = log_file.stat().st_size
        
        # Main loop: tail all active session files for new content
        while True:
            # Discover current session files
            current_logs = sorted(
                [p for p in session_dir.glob("*.jsonl") if ".deleted" not in p.name and ".lock" not in p.name],
                key=lambda p: p.stat().st_mtime, reverse=True
            )[:5]
            
            had_new_data = False
            for log_file in current_logs:
                path_str = str(log_file)
                current_size = log_file.stat().st_size
                last_pos = file_positions.get(path_str, 0)
                
                # New file we haven't seen
                if path_str not in file_positions:
                    file_positions[path_str] = 0
                    last_pos = 0
                
                if current_size > last_pos:
                    with open(log_file, "r") as f:
                        f.seek(last_pos)
                        new_data = f.read()
                        file_positions[path_str] = f.tell()
                    
                    for line in new_data.strip().split("\n"):
                        if line.strip():
                            await _process_activity_line(line, websocket)
                            had_new_data = True
            
            if not had_new_data:
                await asyncio.sleep(0.3)
    except WebSocketDisconnect:
        pass
    except Exception:
        pass


async def _process_activity_line(line: str, websocket: WebSocket):
    """Parse a JSONL line and send tool call / thinking events to the activity feed."""
    try:
        entry = json.loads(line.strip())
        timestamp = entry.get("timestamp", "")
        msg = entry.get("message", {})
        content = msg.get("content", "")
        role = msg.get("role", "")
        
        if isinstance(content, list):
            for block in content:
                btype = block.get("type", "")
                if btype == "toolCall":
                    tool_name = block.get("name", "unknown")
                    tool_input = block.get("arguments", block.get("input", {}))
                    if not isinstance(tool_input, dict):
                        tool_input = {}
                    tl = tool_name.lower()
                    detail = ""
                    code = ""
                    if tl in ("edit", "write"):
                        fp = tool_input.get("file_path", tool_input.get("path", ""))
                        detail = f"FILE: {fp}"
                        if tl == "edit":
                            code = tool_input.get("newText", tool_input.get("new_string", ""))[:500]
                        else:
                            code = tool_input.get("content", "")[:800]
                    elif tl == "exec":
                        detail = tool_input.get("command", "")[:300]
                    elif tl == "read":
                        fp = tool_input.get("file_path", tool_input.get("path", ""))
                        detail = f"FILE: {fp}"
                    elif tl in ("web_search", "web_fetch"):
                        detail = tool_input.get("query", tool_input.get("url", ""))[:200]
                    elif tl == "memory_search":
                        detail = tool_input.get("query", "")[:200]
                    elif tl in ("sessions_spawn", "sessions_send"):
                        detail = tool_input.get("task", tool_input.get("message", ""))[:200]
                    elif tl == "message":
                        detail = tool_input.get("message", tool_input.get("action", ""))[:200]
                    else:
                        detail = json.dumps(tool_input)[:200]
                    
                    await websocket.send_text(json.dumps({
                        "type": "tool_call",
                        "tool": tool_name,
                        "detail": detail,
                        "code": code,
                        "timestamp": timestamp,
                    }))
                
                elif btype == "text" and role == "assistant":
                    text = block.get("text", "")
                    if text.strip():
                        await websocket.send_text(json.dumps({
                            "type": "thinking",
                            "text": text[:300],
                            "timestamp": timestamp,
                        }))
    except (json.JSONDecodeError, KeyError):
        pass


if __name__ == "__main__":
    if not GATEWAY_TOKEN:
        print("⚠️  Set OPENCLAW_GATEWAY_TOKEN environment variable")
    uvicorn.run(app, host="0.0.0.0", port=NERV_PORT, log_level="info")
